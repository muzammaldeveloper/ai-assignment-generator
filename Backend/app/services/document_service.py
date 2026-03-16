"""
Document Generation Service.

Generates professional DOCX and PDF documents from
assignment content with templates, images, and formatting.
"""

from __future__ import annotations

import os
from typing import List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Image as RLImage,
    PageBreak,
)

from app.services.text_generation_service import GeneratedContent, GeneratedSection
from app.services.image_generation_service import GeneratedImage
from app.utils.file_helpers import ensure_directory, generate_unique_filename
from app.utils.logger import get_logger

logger = get_logger(__name__)


# ── Template Color Schemes ──
TEMPLATE_COLORS = {
    "professional": {"heading": RGBColor(0, 51, 102), "accent": RGBColor(0, 102, 153)},
    "academic": {"heading": RGBColor(51, 51, 51), "accent": RGBColor(102, 102, 102)},
    "modern": {"heading": RGBColor(0, 120, 215), "accent": RGBColor(0, 153, 204)},
    "minimal": {"heading": RGBColor(33, 33, 33), "accent": RGBColor(100, 100, 100)},
    "colorful": {"heading": RGBColor(156, 39, 176), "accent": RGBColor(233, 30, 99)},
}


class DocumentService:
    """
    Professional document generation service.

    Creates DOCX and PDF files from generated assignment content
    with proper formatting, images, captions, and references.
    """

    def __init__(self, storage_path: str = "storage/documents") -> None:
        self._storage_path = os.path.abspath(storage_path)
        ensure_directory(self._storage_path)
        logger.info("DocumentService initialized | path=%s", self._storage_path)

    # ══════════════════════════════════════════════════
    #  DOCX Generation
    # ══════════════════════════════════════════════════

    def generate_docx(
        self,
        content: GeneratedContent,
        images: List[GeneratedImage],
        template: str = "professional",
        assignment_id: str = "",
    ) -> str:
        """
        Generate a professional DOCX document.

        Args:
            content: Generated assignment content.
            images: Generated images list.
            template: Template name for color scheme.
            assignment_id: Assignment ID for filename.

        Returns:
            str: Path to the generated DOCX file.
        """
        logger.info("Generating DOCX | title='%s' | template=%s", content.title, template)

        doc = Document()
        colors = TEMPLATE_COLORS.get(template, TEMPLATE_COLORS["professional"])

        # ── Configure default styles ──
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5

        # ── Title ──
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(content.title)
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = colors["heading"]
        doc.add_paragraph("")  # Spacer

        # ── Introduction ──
        self._add_heading(doc, "Introduction", colors["heading"])
        intro_para = doc.add_paragraph(content.introduction)
        intro_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph("")

        # ── Body Sections ──
        image_map = {img.section_title: img for img in images if img.success}

        for section in content.sections:
            self._add_heading(doc, section.title, colors["heading"])

            body_para = doc.add_paragraph(section.content)
            body_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Add image if available
            img = image_map.get(section.title)
            if img and img.image_path and os.path.exists(img.image_path):
                doc.add_paragraph("")
                try:
                    doc.add_picture(img.image_path, width=Inches(5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    caption_para = doc.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.add_run(img.caption)
                    caption_run.italic = True
                    caption_run.font.size = Pt(10)
                    caption_run.font.color.rgb = colors["accent"]
                except Exception as e:
                    logger.warning("Failed to add image to DOCX | error=%s", str(e))

            doc.add_paragraph("")

        # ── Conclusion ──
        self._add_heading(doc, "Conclusion", colors["heading"])
        conclusion_para = doc.add_paragraph(content.conclusion)
        conclusion_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph("")

        # ── References ──
        self._add_heading(doc, "References", colors["heading"])
        for ref in content.references:
            ref_para = doc.add_paragraph(ref)
            ref_para.paragraph_format.space_after = Pt(4)
            ref_para.paragraph_format.first_line_indent = Inches(-0.5)
            ref_para.paragraph_format.left_indent = Inches(0.5)

        # ── Save ──
        filename = generate_unique_filename("docx", prefix=f"assignment_{assignment_id[:8]}")
        filepath = os.path.join(self._storage_path, filename)
        doc.save(filepath)

        logger.info("DOCX generated | path=%s", filepath)
        return filepath

    def _add_heading(self, doc: Document, text: str, color: RGBColor) -> None:
        """Add a styled heading to the document."""
        heading = doc.add_heading(level=1)
        run = heading.add_run(text)
        run.font.size = Pt(18)
        run.font.color.rgb = color
        run.bold = True

    # ══════════════════════════════════════════════════
    #  PDF Generation
    # ══════════════════════════════════════════════════

    def generate_pdf(
        self,
        content: GeneratedContent,
        images: List[GeneratedImage],
        template: str = "professional",
        assignment_id: str = "",
    ) -> str:
        """
        Generate a professional PDF document.

        Args:
            content: Generated assignment content.
            images: Generated images list.
            template: Template name for styling.
            assignment_id: Assignment ID for filename.

        Returns:
            str: Path to the generated PDF file.
        """
        logger.info("Generating PDF | title='%s' | template=%s", content.title, template)

        filename = generate_unique_filename("pdf", prefix=f"assignment_{assignment_id[:8]}")
        filepath = os.path.join(self._storage_path, filename)

        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        # ── Styles ──
        styles = getSampleStyleSheet()
        color_hex = "#003366" if template == "professional" else "#333333"

        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Title"],
            fontSize=24,
            textColor=HexColor(color_hex),
            alignment=1,  # Center
            spaceAfter=30,
        )
        heading_style = ParagraphStyle(
            "CustomHeading",
            parent=styles["Heading1"],
            fontSize=18,
            textColor=HexColor(color_hex),
            spaceBefore=20,
            spaceAfter=10,
        )
        body_style = ParagraphStyle(
            "CustomBody",
            parent=styles["Normal"],
            fontSize=12,
            leading=18,
            alignment=4,  # Justify
            spaceAfter=8,
        )
        caption_style = ParagraphStyle(
            "CustomCaption",
            parent=styles["Normal"],
            fontSize=10,
            textColor=HexColor("#666666"),
            alignment=1,  # Center
            fontName="Helvetica-Oblique",
            spaceAfter=15,
        )

        # ── Build Story ──
        story = []

        # Title
        story.append(Paragraph(content.title, title_style))
        story.append(Spacer(1, 20))

        # Introduction
        story.append(Paragraph("Introduction", heading_style))
        for para in content.introduction.split("\n\n"):
            if para.strip():
                story.append(Paragraph(para.strip(), body_style))
        story.append(Spacer(1, 10))

        # Body Sections
        image_map = {img.section_title: img for img in images if img.success}

        for section in content.sections:
            story.append(Paragraph(section.title, heading_style))

            for para in section.content.split("\n\n"):
                if para.strip():
                    story.append(Paragraph(para.strip(), body_style))

            # Add image if available
            img = image_map.get(section.title)
            if img and img.image_path and os.path.exists(img.image_path):
                try:
                    story.append(Spacer(1, 10))
                    rl_img = RLImage(img.image_path, width=5 * inch, height=3 * inch)
                    rl_img.hAlign = "CENTER"
                    story.append(rl_img)
                    story.append(Paragraph(img.caption, caption_style))
                except Exception as e:
                    logger.warning("Failed to add image to PDF | error=%s", str(e))

            story.append(Spacer(1, 10))

        # Conclusion
        story.append(Paragraph("Conclusion", heading_style))
        for para in content.conclusion.split("\n\n"):
            if para.strip():
                story.append(Paragraph(para.strip(), body_style))
        story.append(Spacer(1, 15))

        # References
        story.append(PageBreak())
        story.append(Paragraph("References", heading_style))
        for ref in content.references:
            story.append(Paragraph(ref, body_style))

        # Build PDF
        doc.build(story)

        logger.info("PDF generated | path=%s", filepath)
        return filepath